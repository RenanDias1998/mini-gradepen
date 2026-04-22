import base64
from datetime import datetime
from io import BytesIO
from urllib.parse import urlencode

import cv2
import numpy as np
import pandas as pd
import qrcode
import streamlit as st
from PIL import Image
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from supabase import Client, create_client


CHOICES = ["A", "B", "C", "D", "E"]
TOTAL_QUESTIONS = 10
TOTAL_ALTERNATIVES = 5
MARKER_CANVAS = (900, 1300)
DEFAULT_GRID = {
    "columns": np.array([218, 354, 490, 626, 762], dtype=np.int32),
    "rows": np.array([282, 376, 469, 563, 657, 751, 844, 938, 1032, 1126], dtype=np.int32),
}
DEFAULT_EXAM_HEADER = (
    "Escola:_____________________________________________________________\n"
    "Nome:_____________________________________________________________\n"
    "Data:_____/_____/_______\n"
    "Serie:________________________________"
)
def init_state():
    defaults = {
        "draft_answer_key": [""] * TOTAL_QUESTIONS,
        "saved_answer_key": [],
        "saved_answer_key_version": 0,
        "last_processed": None,
        "active_exam_code": "",
        "active_exam": None,
        "nav_page": "1. Criar prova",
        "sidebar_page": "1. Criar prova",
        "draft_exam_header": DEFAULT_EXAM_HEADER,
        "draft_objective_texts": [""] * TOTAL_QUESTIONS,
        "draft_objective_options": [[f"Alternativa {choice}" for choice in CHOICES] for _ in range(TOTAL_QUESTIONS)],
        "draft_essay_texts": [],
        "exam_header_data": DEFAULT_EXAM_HEADER,
        "objective_texts_data": [""] * TOTAL_QUESTIONS,
        "objective_options_data": [[f"Alternativa {choice}" for choice in CHOICES] for _ in range(TOTAL_QUESTIONS)],
        "essay_texts_data": [],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def pil_to_bgr(image):
    rgb = np.array(image.convert("RGB"))
    return cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)


def bgr_to_rgb(image):
    return cv2.cvtColor(image, cv2.COLOR_BGR2RGB)


def answer_key_complete(answer_key, total_questions=TOTAL_QUESTIONS):
    return len(answer_key) >= total_questions and all(answer_key[:total_questions])


def get_exam_objective_count(exam):
    return int(exam.get("objective_questions", TOTAL_QUESTIONS))


def get_exam_essay_count(exam):
    return int(exam.get("essay_questions", 0))


def generate_exam_code(title, exam_date):
    base = slugify_exam_code(f"{title}-{exam_date}")
    timestamp = datetime.now().strftime("%H%M")
    return f"{base}-{timestamp}" if base else f"PROVA-{timestamp}"


def go_to_page(page_name):
    st.session_state.nav_page = page_name
    st.session_state.sidebar_page = page_name
    st.rerun()


def ensure_question_drafts(objective_count, essay_count):
    objective_count = int(objective_count)
    essay_count = int(essay_count)

    current_objectives = st.session_state.draft_objective_texts
    if len(current_objectives) < objective_count:
        current_objectives = current_objectives + [""] * (objective_count - len(current_objectives))
    else:
        current_objectives = current_objectives[:objective_count]
    st.session_state.draft_objective_texts = current_objectives

    current_options = st.session_state.draft_objective_options
    default_option_row = [f"Alternativa {choice}" for choice in CHOICES]
    if len(current_options) < objective_count:
        current_options = current_options + [default_option_row.copy() for _ in range(objective_count - len(current_options))]
    else:
        current_options = current_options[:objective_count]
    st.session_state.draft_objective_options = current_options

    current_essays = st.session_state.draft_essay_texts
    if len(current_essays) < essay_count:
        current_essays = current_essays + [""] * (essay_count - len(current_essays))
    else:
        current_essays = current_essays[:essay_count]
    st.session_state.draft_essay_texts = current_essays

    base_objectives = st.session_state.objective_texts_data
    if len(base_objectives) < objective_count:
        base_objectives = base_objectives + [""] * (objective_count - len(base_objectives))
    else:
        base_objectives = base_objectives[:objective_count]
    st.session_state.objective_texts_data = base_objectives

    base_objective_options = st.session_state.objective_options_data
    if len(base_objective_options) < objective_count:
        base_objective_options = base_objective_options + [default_option_row.copy() for _ in range(objective_count - len(base_objective_options))]
    else:
        base_objective_options = base_objective_options[:objective_count]
    st.session_state.objective_options_data = base_objective_options

    base_essays = st.session_state.essay_texts_data
    if len(base_essays) < essay_count:
        base_essays = base_essays + [""] * (essay_count - len(base_essays))
    else:
        base_essays = base_essays[:essay_count]
    st.session_state.essay_texts_data = base_essays


def slugify_exam_code(text):
    allowed = []
    for char in text.strip().upper():
        if char.isalnum():
            allowed.append(char)
        elif char in ("-", "_", "/"):
            allowed.append("-")
        elif char.isspace():
            allowed.append("-")
    slug = "".join(allowed).replace("--", "-").strip("-")
    return slug or "PROVA"


@st.cache_resource
def get_supabase() -> Client:
    url = None
    key = None

    try:
        url = st.secrets["SUPABASE_URL"]
    except Exception:
        pass

    try:
        key = st.secrets["SUPABASE_KEY"]
    except Exception:
        pass

    if (not url or not key) and "supabase" in st.secrets:
        supabase_block = st.secrets["supabase"]
        url = url or supabase_block.get("url")
        key = key or supabase_block.get("key")

    if not url or not key:
        raise RuntimeError(
            "Configure os secrets do Supabase no Streamlit. Use SUPABASE_URL e SUPABASE_KEY "
            "ou o bloco [supabase] com url e key."
        )

    return create_client(url, key)


def map_exam_record(record, answer_key_records=None):
    answers = []
    if answer_key_records:
        sorted_records = sorted(answer_key_records, key=lambda item: item["questao"])
        answers = [item["resposta"] for item in sorted_records]
    elif "gabaritos" in record and record["gabaritos"]:
        sorted_records = sorted(record["gabaritos"], key=lambda item: item["questao"])
        answers = [item["resposta"] for item in sorted_records]

    return {
        "id": record["id"],
        "code": record["codigo"],
        "title": record["titulo"],
        "date": record.get("data_prova") or "",
        "objective_questions": int(record.get("qtd_objetivas", TOTAL_QUESTIONS)),
        "essay_questions": int(record.get("qtd_dissertativas", 0)),
        "answer_key": answers,
        "version": int(record.get("versao", 1)),
        "updated_at": record.get("atualizado_em", ""),
        "header": DEFAULT_EXAM_HEADER,
        "objective_texts": [""] * int(record.get("qtd_objetivas", TOTAL_QUESTIONS)),
        "objective_options": [[f"Alternativa {choice}" for choice in CHOICES] for _ in range(int(record.get("qtd_objetivas", TOTAL_QUESTIONS)))],
        "essay_texts": [""] * int(record.get("qtd_dissertativas", 0)),
    }


def load_exams():
    supabase = get_supabase()
    response = (
        supabase.table("provas")
        .select("id,codigo,titulo,data_prova,qtd_objetivas,qtd_dissertativas,versao,atualizado_em,gabaritos(questao,resposta)")
        .order("titulo")
        .execute()
    )
    exams = {}
    for record in response.data or []:
        exam = map_exam_record(record)
        exams[exam["code"]] = exam
    return exams


def save_exam(exam):
    supabase = get_supabase()
    payload = {
        "codigo": exam["code"],
        "titulo": exam["title"],
        "data_prova": exam["date"] or None,
        "qtd_objetivas": exam["objective_questions"],
        "qtd_dissertativas": exam["essay_questions"],
        "versao": exam.get("version", 1),
        "atualizado_em": datetime.now().isoformat(),
    }
    response = (
        supabase.table("provas")
        .upsert(payload, on_conflict="codigo")
        .execute()
    )
    record = response.data[0]

    supabase.table("gabaritos").delete().eq("prova_id", record["id"]).execute()
    answer_key = exam.get("answer_key", [])
    if answer_key:
        rows = [
            {
                "prova_id": record["id"],
                "questao": index + 1,
                "resposta": answer,
            }
            for index, answer in enumerate(answer_key)
        ]
        supabase.table("gabaritos").insert(rows).execute()

    fresh_exam = (
        supabase.table("provas")
        .select("id,codigo,titulo,data_prova,qtd_objetivas,qtd_dissertativas,versao,atualizado_em,gabaritos(questao,resposta)")
        .eq("codigo", exam["code"])
        .single()
        .execute()
    )
    saved_exam = map_exam_record(fresh_exam.data)
    for optional_key in ("header", "objective_texts", "objective_options", "essay_texts"):
        if optional_key in exam:
            saved_exam[optional_key] = exam[optional_key]
    return saved_exam


def load_results(exam_code=None):
    supabase = get_supabase()
    query = (
        supabase.table("resultados")
        .select("id,aluno,turma,acertos,erros,respostas_lidas,questoes_certas,questoes_erradas,criado_em,provas(codigo,titulo,data_prova)")
        .order("aluno")
    )
    if exam_code:
        query = query.eq("provas.codigo", exam_code)
    response = query.execute()

    records = []
    for item in response.data or []:
        prova = item.get("provas") or {}
        records.append(
            {
                "Prova": prova.get("titulo", ""),
                "Codigo Prova": prova.get("codigo", ""),
                "Data Prova": prova.get("data_prova", ""),
                "Aluno": item.get("aluno", ""),
                "Turma": item.get("turma", ""),
                "Acertos": item.get("acertos", 0),
                "Erros": item.get("erros", 0),
                "Questoes Certas": item.get("questoes_certas", ""),
                "Questoes Erradas": item.get("questoes_erradas", ""),
                "Respostas Lidas": item.get("respostas_lidas", ""),
            }
        )
    return records


def save_result(record, exam_id):
    supabase = get_supabase()
    payload = {
        "prova_id": exam_id,
        "aluno": record["Aluno"],
        "turma": record["Turma"],
        "acertos": record["Acertos"],
        "erros": record["Erros"],
        "respostas_lidas": record["Respostas Lidas"],
        "questoes_certas": record["Questoes Certas"],
        "questoes_erradas": record["Questoes Erradas"],
    }
    supabase.table("resultados").insert(payload).execute()


def build_exam_url(base_url, exam_code):
    clean_base = base_url.strip().rstrip("/")
    if not clean_base:
        return ""
    return f"{clean_base}?{urlencode({'prova': exam_code})}"


def build_qr_image(data):
    qr = qrcode.QRCode(version=2, box_size=8, border=2)
    qr.add_data(data)
    qr.make(fit=True)
    return qr.make_image(fill_color="black", back_color="white").convert("RGB")


def qr_image_to_base64(image):
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def set_active_exam(exam):
    st.session_state.active_exam = exam
    st.session_state.active_exam_code = exam["code"]
    st.session_state.saved_answer_key = exam["answer_key"]
    st.session_state.saved_answer_key_version = exam.get("version", 1)
    st.session_state.draft_answer_key = exam["answer_key"].copy()
    st.session_state.exam_header_data = exam.get("header", "")
    st.session_state.objective_texts_data = exam.get("objective_texts", []).copy()
    st.session_state.objective_options_data = exam.get("objective_options", []).copy()
    st.session_state.essay_texts_data = exam.get("essay_texts", []).copy()


def sync_exam_from_query(exams):
    query_params = st.query_params
    exam_code = query_params.get("prova", "")
    if not exam_code:
        return

    exam_code = slugify_exam_code(exam_code)
    exam = exams.get(exam_code)
    if exam is None:
        st.warning(f"O link da prova '{exam_code}' nao foi encontrado no cadastro local.")
        return

    if st.session_state.active_exam_code != exam_code:
        set_active_exam(exam)


def order_points(points):
    pts = np.array(points, dtype=np.float32)
    sums = pts.sum(axis=1)
    diffs = np.diff(pts, axis=1).reshape(-1)
    return np.array(
        [
            pts[np.argmin(sums)],
            pts[np.argmin(diffs)],
            pts[np.argmax(sums)],
            pts[np.argmax(diffs)],
        ],
        dtype=np.float32,
    )


def four_point_transform(image, points, size=None):
    rect = order_points(points)
    (tl, tr, br, bl) = rect

    if size is None:
        width_a = np.linalg.norm(br - bl)
        width_b = np.linalg.norm(tr - tl)
        height_a = np.linalg.norm(tr - br)
        height_b = np.linalg.norm(tl - bl)
        max_width = max(int(width_a), int(width_b))
        max_height = max(int(height_a), int(height_b))
    else:
        max_width, max_height = size

    destination = np.array(
        [
            [0, 0],
            [max_width - 1, 0],
            [max_width - 1, max_height - 1],
            [0, max_height - 1],
        ],
        dtype=np.float32,
    )

    matrix = cv2.getPerspectiveTransform(rect, destination)
    warped = cv2.warpPerspective(image, matrix, (max_width, max_height))
    return warped, rect


def resize_for_preview(image, max_width=1100):
    height, width = image.shape[:2]
    if width <= max_width:
        return image.copy(), 1.0
    scale = max_width / width
    resized = cv2.resize(image, None, fx=scale, fy=scale, interpolation=cv2.INTER_AREA)
    return resized, scale


def preprocess_gray(gray):
    clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(8, 8))
    enhanced = clahe.apply(gray)
    blurred = cv2.GaussianBlur(enhanced, (5, 5), 0)
    return enhanced, blurred


def detect_sheet(image):
    preview, scale = resize_for_preview(image)
    gray = cv2.cvtColor(preview, cv2.COLOR_BGR2GRAY)
    _, blurred = preprocess_gray(gray)
    edges = cv2.Canny(blurred, 50, 150)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
    edges = cv2.dilate(edges, kernel, iterations=2)
    edges = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel, iterations=2)

    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    contours = sorted(contours, key=cv2.contourArea, reverse=True)

    chosen = None
    for contour in contours[:15]:
        perimeter = cv2.arcLength(contour, True)
        approx = cv2.approxPolyDP(contour, 0.02 * perimeter, True)
        area = cv2.contourArea(contour)
        if len(approx) == 4 and area > preview.shape[0] * preview.shape[1] * 0.15:
            chosen = approx.reshape(4, 2).astype(np.float32) / scale
            break

    if chosen is None and contours:
        rect = cv2.minAreaRect(contours[0])
        chosen = cv2.boxPoints(rect).astype(np.float32) / scale

    if chosen is None:
        raise ValueError("Nao foi possivel detectar a folha automaticamente.")

    warped, ordered = four_point_transform(image, chosen)
    preview_outline = preview.copy()
    preview_points = (order_points(chosen) * scale).astype(np.int32)
    cv2.polylines(preview_outline, [preview_points], True, (0, 255, 0), 4)

    return warped, {"folha_detectada": preview_outline, "bordas_folha": edges, "pontos_folha": ordered}


def threshold_for_dark_regions(gray):
    enhanced, blurred = preprocess_gray(gray)
    adaptive = cv2.adaptiveThreshold(
        blurred,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY_INV,
        35,
        11,
    )
    _, otsu = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
    merged = cv2.bitwise_or(adaptive, otsu)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
    merged = cv2.morphologyEx(merged, cv2.MORPH_OPEN, kernel, iterations=1)
    merged = cv2.morphologyEx(merged, cv2.MORPH_CLOSE, kernel, iterations=1)
    return enhanced, merged


def sort_markers(markers):
    ordered = sorted(markers, key=lambda marker: marker["center"][1])
    top = sorted(ordered[:2], key=lambda marker: marker["center"][0])
    bottom = sorted(ordered[2:], key=lambda marker: marker["center"][0])
    return [top[0], top[1], bottom[1], bottom[0]]


def find_reference_markers(sheet):
    gray = cv2.cvtColor(sheet, cv2.COLOR_BGR2GRAY)
    enhanced, dark = threshold_for_dark_regions(gray)
    height, width = dark.shape
    search = np.zeros_like(dark)
    search[:, int(width * 0.48):] = dark[:, int(width * 0.48):]

    contours, _ = cv2.findContours(search, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    candidates = []

    for contour in contours:
        area = cv2.contourArea(contour)
        if area < 200 or area > height * width * 0.04:
            continue

        x, y, w, h = cv2.boundingRect(contour)
        aspect = w / float(h)
        if not 0.7 <= aspect <= 1.3:
            continue

        fill_ratio = area / float(w * h)
        if fill_ratio < 0.65:
            continue

        perimeter = cv2.arcLength(contour, True)
        approx = cv2.approxPolyDP(contour, 0.04 * perimeter, True)
        if len(approx) < 4:
            continue

        candidates.append(
            {
                "rect": (x, y, w, h),
                "center": (x + w / 2.0, y + h / 2.0),
                "area": area,
            }
        )

    if len(candidates) < 4:
        raise ValueError("Os quadrados pretos de referencia nao foram encontrados.")

    candidates.sort(key=lambda item: item["area"], reverse=True)
    best_group = None

    for start in range(len(candidates)):
        for end in range(start + 4, len(candidates) + 1):
            group = candidates[start:end]
            if len(group) < 4:
                continue

            centers = np.array([item["center"] for item in group], dtype=np.float32)
            median_x = np.median(centers[:, 0])
            top = centers[centers[:, 1] <= np.median(centers[:, 1])]
            bottom = centers[centers[:, 1] > np.median(centers[:, 1])]
            left = centers[centers[:, 0] <= median_x]
            right = centers[centers[:, 0] > median_x]

            if len(top) >= 2 and len(bottom) >= 2 and len(left) >= 2 and len(right) >= 2:
                best_group = sorted(group, key=lambda item: item["area"], reverse=True)[:4]
                break
        if best_group:
            break

    if best_group is None:
        best_group = candidates[:4]

    ordered_markers = sort_markers(best_group)
    tl, tr, br, bl = ordered_markers
    marker_points = np.array(
        [
            [tl["rect"][0], tl["rect"][1]],
            [tr["rect"][0] + tr["rect"][2], tr["rect"][1]],
            [br["rect"][0] + br["rect"][2], br["rect"][1] + br["rect"][3]],
            [bl["rect"][0], bl["rect"][1] + bl["rect"][3]],
        ],
        dtype=np.float32,
    )
    marker_debug = cv2.cvtColor(enhanced, cv2.COLOR_GRAY2BGR)

    for index, marker in enumerate(best_group, start=1):
        x, y, w, h = marker["rect"]
        cv2.rectangle(marker_debug, (x, y), (x + w, y + h), (0, 255, 0), 2)
        cx, cy = map(int, marker["center"])
        cv2.putText(
            marker_debug,
            str(index),
            (cx - 10, cy - 10),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.8,
            (0, 0, 255),
            2,
            cv2.LINE_AA,
        )

    return marker_points, {"marcadores": marker_debug, "mascara_escura": dark}


def infer_positions(binary):
    height, width = binary.shape

    col_window = binary[int(height * 0.18): int(height * 0.88), int(width * 0.15): int(width * 0.92)]
    vertical_projection = col_window.sum(axis=0).astype(np.float32)
    col_indices = np.where(vertical_projection > vertical_projection.max() * 0.45)[0]

    if len(col_indices) > 0:
        groups = np.split(col_indices, np.where(np.diff(col_indices) > 12)[0] + 1)
        centers = np.array(
            [int(group.mean()) + int(width * 0.15) for group in groups if len(group) > 4],
            dtype=np.int32,
        )
    else:
        centers = np.array([], dtype=np.int32)

    if len(centers) != TOTAL_ALTERNATIVES:
        centers = DEFAULT_GRID["columns"]

    row_window = binary[int(height * 0.18): int(height * 0.92), int(width * 0.18): int(width * 0.86)]
    horizontal_projection = row_window.sum(axis=1).astype(np.float32)
    row_indices = np.where(horizontal_projection > horizontal_projection.max() * 0.55)[0]

    if len(row_indices) > 0:
        groups = np.split(row_indices, np.where(np.diff(row_indices) > 14)[0] + 1)
        rows = np.array(
            [int(group.mean()) + int(height * 0.18) for group in groups if len(group) > 4],
            dtype=np.int32,
        )
    else:
        rows = np.array([], dtype=np.int32)

    if len(rows) != TOTAL_QUESTIONS:
        rows = DEFAULT_GRID["rows"]

    return centers, rows


def score_bubbles(warped_answers):
    gray = cv2.cvtColor(warped_answers, cv2.COLOR_BGR2GRAY)
    normalized = cv2.normalize(gray, None, 0, 255, cv2.NORM_MINMAX)
    adaptive = cv2.adaptiveThreshold(
        normalized,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY_INV,
        31,
        9,
    )
    kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    cleaned = cv2.morphologyEx(adaptive, cv2.MORPH_OPEN, kernel, iterations=1)
    cleaned = cv2.morphologyEx(cleaned, cv2.MORPH_CLOSE, kernel, iterations=1)

    columns, rows = infer_positions(cleaned)
    radius = 28
    scores = []
    answers = []

    for center_y in rows:
        row_scores = []
        for center_x in columns:
            mask = np.zeros_like(cleaned)
            cv2.circle(mask, (int(center_x), int(center_y)), radius, 255, -1)
            focused = cv2.bitwise_and(cleaned, cleaned, mask=mask)
            dark_pixels = cv2.countNonZero(focused)
            area = cv2.countNonZero(mask)
            row_scores.append(dark_pixels / float(max(area, 1)))

        scores.append(row_scores)
        best_index = int(np.argmax(row_scores))
        sorted_scores = np.sort(np.array(row_scores))[::-1]
        confident = sorted_scores[0] > 0.18 and (sorted_scores[0] - sorted_scores[1] > 0.035)
        answers.append(CHOICES[best_index] if confident else "?")

    diagnostics = {
        "binaria_bolhas": cleaned,
        "colunas": columns,
        "linhas": rows,
        "scores": scores,
    }
    return answers, diagnostics


def build_answer_overlay(warped_answers, answers, diagnostics):
    overlay = warped_answers.copy()
    columns = diagnostics["colunas"]
    rows = diagnostics["linhas"]
    scores = diagnostics["scores"]

    for question_index, center_y in enumerate(rows):
        cv2.putText(
            overlay,
            f"Q{question_index + 1}",
            (28, int(center_y + 8)),
            cv2.FONT_HERSHEY_SIMPLEX,
            0.75,
            (40, 40, 40),
            2,
            cv2.LINE_AA,
        )
        for choice_index, center_x in enumerate(columns):
            selected = answers[question_index] == CHOICES[choice_index]
            color = (0, 180, 0) if selected else (0, 140, 255)
            cv2.circle(overlay, (int(center_x), int(center_y)), 30, color, 2)
            cv2.putText(
                overlay,
                f"{scores[question_index][choice_index]:.2f}",
                (int(center_x - 28), int(center_y - 38)),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.42,
                color,
                1,
                cv2.LINE_AA,
            )

    return overlay


def extract_answer_area(sheet):
    marker_points, marker_diagnostics = find_reference_markers(sheet)
    warped_answers, _ = four_point_transform(sheet, marker_points, size=MARKER_CANVAS)
    preview = sheet.copy()
    cv2.polylines(preview, [marker_points.astype(np.int32)], True, (255, 0, 0), 4)

    return warped_answers, {
        "area_gabarito": preview,
        **marker_diagnostics,
    }


def process_answer_sheet(image_bgr):
    sheet, sheet_diagnostics = detect_sheet(image_bgr)
    answer_area, marker_diagnostics = extract_answer_area(sheet)
    answers, bubble_diagnostics = score_bubbles(answer_area)
    overlay = build_answer_overlay(answer_area, answers, bubble_diagnostics)

    return {
        "folha": sheet,
        "gabarito": answer_area,
        "respostas": answers,
        "diagnosticos": {
            **sheet_diagnostics,
            **marker_diagnostics,
            **bubble_diagnostics,
            "overlay_leitura": overlay,
        },
    }


def render_diagnostics(result):
    st.subheader("Diagnostico visual")
    col1, col2 = st.columns(2)
    with col1:
        st.image(bgr_to_rgb(result["diagnosticos"]["folha_detectada"]), caption="Folha detectada")
        st.image(result["diagnosticos"]["bordas_folha"], caption="Bordas usadas na deteccao")
        st.image(bgr_to_rgb(result["diagnosticos"]["marcadores"]), caption="Quadrados pretos encontrados")
    with col2:
        st.image(bgr_to_rgb(result["diagnosticos"]["area_gabarito"]), caption="Recorte bruto do bloco de respostas")
        st.image(bgr_to_rgb(result["gabarito"]), caption="Gabarito alinhado")
        st.image(bgr_to_rgb(result["diagnosticos"]["overlay_leitura"]), caption="Leitura com scores por bolha")

    st.image(result["diagnosticos"]["binaria_bolhas"], caption="Mascara binaria usada na leitura")


def select_answer(question_index, choice):
    st.session_state.draft_answer_key[question_index] = choice


def render_answer_key_editor(total_questions):
    st.subheader("Gabarito do professor")
    st.caption("Clique na alternativa de cada questao. A escolha permanece ate voce salvar.")

    for question_index in range(total_questions):
        current_choice = st.session_state.draft_answer_key[question_index]
        row_columns = st.columns([1.2, 1, 1, 1, 1, 1, 1.6])
        row_columns[0].markdown(f"**Q{question_index + 1}**")

        for choice_index, choice in enumerate(CHOICES, start=1):
            selected = current_choice == choice
            button_type = "primary" if selected else "secondary"
            row_columns[choice_index].button(
                choice,
                key=f"draft_q{question_index}_{choice}",
                type=button_type,
                use_container_width=True,
                on_click=select_answer,
                args=(question_index, choice),
            )

        if current_choice:
            row_columns[6].markdown(
                f"Selecionada: <span style='color:#15803d; font-weight:700;'>{current_choice}</span>",
                unsafe_allow_html=True,
            )
        else:
            row_columns[6].markdown(
                "<span style='color:#b45309; font-weight:600;'>Aguardando escolha</span>",
                unsafe_allow_html=True,
            )


def compute_score(student_answers, saved_answer_key):
    comparisons = []
    correct = 0

    for question_index, official_answer in enumerate(saved_answer_key):
        student_answer = student_answers[question_index] if question_index < len(student_answers) else "?"
        is_correct = student_answer == official_answer
        correct += int(is_correct)
        comparisons.append(
            {
                "questao": question_index + 1,
                "gabarito": official_answer,
                "aluno": student_answer,
                "status": "Acerto" if is_correct else "Erro",
            }
        )

    return correct, comparisons


def build_student_record(name, class_name, exam, student_answers, correct_answers, comparisons):
    total_questions = len(exam["answer_key"])
    correct_questions = [f"Q{item['questao']}" for item in comparisons if item["status"] == "Acerto"]
    wrong_questions = [f"Q{item['questao']}" for item in comparisons if item["status"] == "Erro"]
    record = {
        "Prova": exam["title"],
        "Codigo Prova": exam["code"],
        "Data Prova": exam.get("date", ""),
        "Aluno": name or "Nao informado",
        "Turma": class_name or "Nao informada",
        "Acertos": correct_answers,
        "Erros": total_questions - correct_answers,
        "Respostas Lidas": " ".join(student_answers),
        "Gabarito Usado": " ".join(exam["answer_key"]),
        "Versao Gabarito": exam.get("version", 1),
        "Questoes Certas": ", ".join(correct_questions),
        "Questoes Erradas": ", ".join(wrong_questions),
    }

    for item in comparisons:
        record[f"Q{item['questao']}"] = item["aluno"]
        record[f"Q{item['questao']}_status"] = item["status"]

    return record


def export_results_to_excel(dataframe=None):
    if dataframe is None:
        if not st.session_state.class_results:
            return None
        dataframe = pd.DataFrame(st.session_state.class_results)

    if dataframe.empty:
        return None
    ordered_columns = [
        "Prova",
        "Codigo Prova",
        "Data Prova",
        "Aluno",
        "Turma",
        "Acertos",
        "Erros",
        "Questoes Certas",
        "Questoes Erradas",
        "Respostas Lidas",
        "Gabarito Usado",
        "Versao Gabarito",
    ]

    for question_index in range(1, TOTAL_QUESTIONS + 1):
        ordered_columns.append(f"Q{question_index}")
        ordered_columns.append(f"Q{question_index}_status")

    data = dataframe[[column for column in ordered_columns if column in dataframe.columns]]

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        data.to_excel(writer, index=False, sheet_name="Correcoes")

    output.seek(0)
    return output


def render_saved_results(active_exam_code=None):
    st.subheader("Turma corrigida")
    records = load_results(active_exam_code)
    if not records:
        st.info("Nenhum aluno foi salvo ainda.")
        return

    dataframe = pd.DataFrame(records)

    if dataframe.empty:
        st.info("Nenhum aluno salvo para esta prova ainda.")
        return

    dataframe = dataframe.sort_values(["Aluno", "Turma"], kind="stable").reset_index(drop=True)
    st.dataframe(dataframe, use_container_width=True)

    excel_file = export_results_to_excel(dataframe)
    if excel_file is not None:
        st.download_button(
            "Baixar Excel da turma",
            data=excel_file,
            file_name="correcoes_turma.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )


def generate_printable_template(exam, exam_url):
    qr_b64 = ""
    if exam_url:
        qr_b64 = qr_image_to_base64(build_qr_image(exam_url))

    qr_html = (
        f'<img src="data:image/png;base64,{qr_b64}" alt="QR da prova" style="width:36mm;height:36mm;display:block;" />'
        if qr_b64
        else '<div style="font-size:9px;color:#666;text-align:center;">Defina a URL publica do app para gerar o QR.</div>'
    )

    rows_html = []
    objective_questions = get_exam_objective_count(exam)
    essay_questions = get_exam_essay_count(exam)
    for question_index in range(objective_questions):
        option_labels = exam.get("objective_options", [])
        current_options = option_labels[question_index] if question_index < len(option_labels) else [f"Alternativa {choice}" for choice in CHOICES]
        bubbles = "".join(
            f'<div style="text-align:center;"><div class="bubble"></div><div style="font-size:8px; margin-top:1mm;">{current_options[idx]}</div></div>'
            for idx, _ in enumerate(CHOICES)
        )
        rows_html.append(
            f'<div class="answer-row"><div class="question-label">Q{question_index + 1}</div>{bubbles}</div>'
        )

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="UTF-8">
  <title>{exam["title"]}</title>
  <style>
    :root {{
      --page-width: 210mm;
      --page-height: 297mm;
      --ink: #111111;
      --soft-ink: #444444;
      --line: #bfc5cc;
      --paper: #ffffff;
      --panel: #f6f7f8;
      --marker-size: 8mm;
      --bubble-size: 5.6mm;
      --row-gap: 3.5mm;
    }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; background: #eceff1; font-family: "Segoe UI", Arial, sans-serif; color: var(--ink); }}
    .page {{ width: var(--page-width); min-height: var(--page-height); margin: 12px auto; background: var(--paper); padding: 8mm 8mm 9mm; }}
    .header {{ border: 1px solid var(--line); padding: 3.5mm 4.5mm; margin-bottom: 4mm; }}
    .header-top {{ display:flex; justify-content:space-between; gap:12px; margin-bottom:2.5mm; }}
    .title {{ font-size:18px; font-weight:700; }}
    .subtitle {{ font-size:10px; color:var(--soft-ink); }}
    .meta-grid {{ display:grid; grid-template-columns:1.2fr 0.9fr 0.9fr; gap:3mm; }}
    .field {{ border:1px solid var(--line); padding:2.2mm; min-height:12mm; }}
    .field-label {{ font-size:10px; font-weight:700; text-transform:uppercase; color:var(--soft-ink); margin-bottom:1.2mm; }}
    .field-line {{ border-bottom:1px solid #949aa1; height:5mm; }}
    .content {{ display:grid; grid-template-columns:58mm 1fr; gap:5mm; align-items:start; }}
    .left-panel, .right-panel {{ border:1px solid var(--line); min-height:170mm; position:relative; }}
    .left-panel {{ padding:4mm; background:linear-gradient(var(--panel), var(--panel)) top/100% 14mm no-repeat, var(--paper); }}
    .panel-title {{ font-size:12px; font-weight:700; text-transform:uppercase; color:var(--soft-ink); margin-bottom:3mm; }}
    .qr-box {{ width:42mm; height:42mm; border:1px solid var(--ink); margin:0 auto 3mm; display:grid; place-items:center; background:#fff; }}
    .id-box {{ border:1px solid var(--line); padding:2.5mm; margin-top:3mm; min-height:16mm; }}
    .instructions {{ margin-top:4mm; padding-left:5mm; font-size:9.5px; line-height:1.3; }}
    .instructions li + li {{ margin-top:1.2mm; }}
    .right-panel {{ padding:6mm 6mm 7mm; }}
    .marker {{ position:absolute; width:var(--marker-size); height:var(--marker-size); background:#000; }}
    .marker.tl {{ top:5mm; left:5mm; }}
    .marker.tr {{ top:5mm; right:5mm; }}
    .marker.bl {{ bottom:5mm; left:5mm; }}
    .marker.br {{ bottom:5mm; right:5mm; }}
    .answer-frame {{ border:1px solid #d1d6db; min-height:155mm; padding:7mm 6mm 6mm; }}
    .answer-guide {{ text-align:center; font-size:10px; color:var(--soft-ink); margin-bottom:4mm; font-weight:600; }}
    .answer-header, .answer-row {{ display:grid; grid-template-columns:15mm repeat(5, 1fr); align-items:center; column-gap:2mm; }}
    .answer-header {{ font-size:11px; font-weight:700; text-align:center; margin-bottom:2.5mm; padding:0 2mm; }}
    .answer-row {{ min-height:8mm; margin-bottom:var(--row-gap); padding:0 2mm; }}
    .question-label {{ font-size:11px; font-weight:700; }}
    .bubble {{ width:var(--bubble-size); height:var(--bubble-size); border:1.2px solid #000; border-radius:50%; margin:0 auto; background:#fff; }}
    .footer {{ margin-top:4mm; border:1px solid var(--line); padding:3mm 4mm; font-size:9.5px; color:var(--soft-ink); display:flex; justify-content:space-between; gap:8mm; flex-wrap:wrap; }}
    @page {{ size:A4 portrait; margin:0; }}
  </style>
</head>
<body>
  <main class="page">
    <section class="header">
      <div class="header-top">
        <div class="title">{exam["title"]}</div>
        <div class="subtitle">Codigo: {exam["code"]} | Data: {exam.get("date", "") or "-"} | Objetivas: {objective_questions} | Dissertativas: {essay_questions}</div>
      </div>
      <div class="meta-grid">
        <div class="field"><div class="field-label">Nome do aluno</div><div class="field-line"></div></div>
        <div class="field"><div class="field-label">Turma</div><div class="field-line"></div></div>
        <div class="field"><div class="field-label">Codigo da prova</div><div class="field-line"></div></div>
      </div>
    </section>
    <section class="content">
      <aside class="left-panel">
        <div class="panel-title">Identificacao</div>
        <div class="qr-box">{qr_html}</div>
        <div class="id-box">
          <div class="field-label">Link da prova</div>
          <div style="font-size:8px; word-break:break-word;">{exam_url or "Defina a URL publica do app para gerar o link."}</div>
        </div>
        <ol class="instructions">
          <li>Preencha somente uma alternativa por questao.</li>
          <li>Pinte a bolha completamente com caneta preta ou azul escura.</li>
          <li>Nao dobre a folha e nao escreva dentro da area do gabarito.</li>
          <li>Mantenha os quadrados pretos livres para alinhamento automatico.</li>
        </ol>
      </aside>
      <section class="right-panel">
        <div class="marker tl"></div>
        <div class="marker tr"></div>
        <div class="marker bl"></div>
        <div class="marker br"></div>
        <div class="answer-frame">
          <div class="answer-guide">Marque o gabarito preenchendo completamente a bolha escolhida</div>
          <div class="answer-header"><div></div><div>A</div><div>B</div><div>C</div><div>D</div><div>E</div></div>
          {"".join(rows_html)}
        </div>
      </section>
    </section>
    <footer class="footer">
      <div>Gabarito vinculado a prova cadastrada no site.</div>
      <div>Objetivas: {objective_questions} | Dissertativas: {essay_questions}</div>
    </footer>
  </main>
</body>
</html>"""


def build_exam_docx(exam):
    document = Document()
    document.add_heading(exam["title"], level=1)

    for line in (exam.get("header") or DEFAULT_EXAM_HEADER).splitlines():
        document.add_paragraph(line)

    document.add_paragraph(f"Codigo da prova: {exam['code']}")
    if exam.get("date"):
        document.add_paragraph(f"Data da prova: {exam['date']}")

    document.add_paragraph("")
    document.add_heading("Questoes objetivas", level=2)
    for index, text in enumerate(exam.get("objective_texts", []), start=1):
        document.add_paragraph(f"Q{index}. {text or '[Enunciado da questao objetiva]'}")
        document.add_paragraph("   ( ) A    ( ) B    ( ) C    ( ) D    ( ) E")

    essay_texts = exam.get("essay_texts", [])
    if essay_texts:
        document.add_heading("Questoes dissertativas", level=2)
        for index, text in enumerate(essay_texts, start=1):
            document.add_paragraph(f"D{index}. {text or '[Enunciado da questao dissertativa]'}")
            document.add_paragraph("_" * 90)
            document.add_paragraph("_" * 90)
            document.add_paragraph("_" * 90)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def draw_wrapped_text(pdf, text, x, y, max_width, font_name="Helvetica", font_size=11, leading=15):
    words = text.split()
    current_line = ""
    lines = []

    for word in words:
        test_line = f"{current_line} {word}".strip()
        if stringWidth(test_line, font_name, font_size) <= max_width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = word

    if current_line:
        lines.append(current_line)

    pdf.setFont(font_name, font_size)
    for line in lines:
        pdf.drawString(x, y, line)
        y -= leading
    return y


def build_exam_pdf(exam):
    output = BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    x_margin = 40
    y = height - 50

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(x_margin, y, exam["title"])
    y -= 28

    pdf.setFont("Helvetica", 11)
    for line in (exam.get("header") or DEFAULT_EXAM_HEADER).splitlines():
        pdf.drawString(x_margin, y, line)
        y -= 16

    pdf.drawString(x_margin, y, f"Codigo da prova: {exam['code']}")
    y -= 16
    if exam.get("date"):
        pdf.drawString(x_margin, y, f"Data da prova: {exam['date']}")
        y -= 22

    pdf.setFont("Helvetica-Bold", 13)
    pdf.drawString(x_margin, y, "Questoes objetivas")
    y -= 20

    for index, text in enumerate(exam.get("objective_texts", []), start=1):
        if y < 90:
            pdf.showPage()
            y = height - 50
        pdf.setFont("Helvetica-Bold", 11)
        pdf.drawString(x_margin, y, f"Q{index}.")
        y = draw_wrapped_text(
            pdf,
            text or "[Enunciado da questao objetiva]",
            x_margin + 28,
            y,
            width - (x_margin * 2) - 28,
        )
        pdf.setFont("Helvetica", 11)
        pdf.drawString(x_margin + 28, y, "( ) A    ( ) B    ( ) C    ( ) D    ( ) E")
        y -= 24

    essay_texts = exam.get("essay_texts", [])
    if essay_texts:
        if y < 120:
            pdf.showPage()
            y = height - 50
        pdf.setFont("Helvetica-Bold", 13)
        pdf.drawString(x_margin, y, "Questoes dissertativas")
        y -= 20

        for index, text in enumerate(essay_texts, start=1):
            if y < 130:
                pdf.showPage()
                y = height - 50
            pdf.setFont("Helvetica-Bold", 11)
            pdf.drawString(x_margin, y, f"D{index}.")
            y = draw_wrapped_text(
                pdf,
                text or "[Enunciado da questao dissertativa]",
                x_margin + 28,
                y,
                width - (x_margin * 2) - 28,
            )
            for _ in range(4):
                pdf.line(x_margin, y, width - x_margin, y)
                y -= 18
            y -= 6

    pdf.save()
    output.seek(0)
    return output


def render_exam_selector(exams, label="Selecione uma prova"):
    if not exams:
        st.info("Nenhuma prova cadastrada ainda.")
        return None

    codes = sorted(exams.keys())
    selected_code = st.selectbox(
        label,
        options=[""] + codes,
        format_func=lambda code: "Escolha uma prova" if not code else f"{code} - {exams[code]['title']}",
        key=f"selector_{label}",
    )
    if selected_code:
        return exams[selected_code]
    return None


def render_create_exam_page():
    st.subheader("1. Criar prova")
    st.caption("Cadastre a estrutura da prova. A leitura automatica cobre as objetivas; as dissertativas ficam registradas no cadastro.")

    exam_title = st.text_input("Titulo da prova", key="exam_title_input")
    exam_date = st.text_input("Data da prova", key="exam_date_input", value=datetime.now().strftime("%Y-%m-%d"))
    objective_questions = st.number_input(
        "Quantidade de questoes objetivas",
        min_value=1,
        max_value=TOTAL_QUESTIONS,
        value=min(10, TOTAL_QUESTIONS),
        step=1,
        key="objective_questions_input",
    )
    essay_questions = st.number_input(
        "Quantidade de questoes dissertativas",
        min_value=0,
        max_value=20,
        value=0,
        step=1,
        key="essay_questions_input",
    )
    base_url = st.text_input(
        "URL publica do app (opcional)",
        key="base_url_input",
        help="No Streamlit Cloud, e algo como https://nome-do-app.streamlit.app",
    )
    ensure_question_drafts(objective_questions, essay_questions)
    if not st.session_state.draft_exam_header and st.session_state.exam_header_data:
        st.session_state.draft_exam_header = st.session_state.exam_header_data
    for question_index in range(int(objective_questions)):
        widget_key = f"objective_text_{question_index}"
        if widget_key not in st.session_state and question_index < len(st.session_state.objective_texts_data):
            st.session_state[widget_key] = st.session_state.objective_texts_data[question_index]
        for choice_index, choice in enumerate(CHOICES):
            option_key = f"objective_option_{question_index}_{choice}"
            if option_key not in st.session_state and question_index < len(st.session_state.objective_options_data):
                options_row = st.session_state.objective_options_data[question_index]
                if choice_index < len(options_row):
                    st.session_state[option_key] = options_row[choice_index]
    for question_index in range(int(essay_questions)):
        widget_key = f"essay_text_{question_index}"
        if widget_key not in st.session_state and question_index < len(st.session_state.essay_texts_data):
            st.session_state[widget_key] = st.session_state.essay_texts_data[question_index]

    st.markdown("**Cabecalho da prova**")
    exam_header = st.text_area(
        "Texto do cabecalho/instrucoes",
        key="draft_exam_header",
        placeholder=DEFAULT_EXAM_HEADER,
        height=140,
    )

    st.markdown("**Questoes objetivas**")
    for question_index in range(int(objective_questions)):
        with st.container(border=True):
            st.markdown(f"**Questao objetiva Q{question_index + 1}**")
            st.session_state.draft_objective_texts[question_index] = st.text_area(
                f"Enunciado da objetiva Q{question_index + 1}",
                key=f"objective_text_{question_index}",
                height=80,
                label_visibility="collapsed",
                placeholder=f"Digite aqui o enunciado da questao objetiva {question_index + 1}",
            )
            option_cols = st.columns(len(CHOICES))
            for choice_index, choice in enumerate(CHOICES):
                st.session_state.draft_objective_options[question_index][choice_index] = option_cols[choice_index].text_input(
                    f"Alternativa {choice} da Q{question_index + 1}",
                    key=f"objective_option_{question_index}_{choice}",
                    placeholder=f"Texto da alternativa {choice}",
                )

    if int(essay_questions) > 0:
        st.markdown("**Questoes dissertativas**")
        for question_index in range(int(essay_questions)):
            st.session_state.draft_essay_texts[question_index] = st.text_area(
                f"Enunciado da dissertativa D{question_index + 1}",
                key=f"essay_text_{question_index}",
                height=90,
                placeholder=f"Digite aqui o enunciado da questao dissertativa {question_index + 1}",
            )

    generated_code = generate_exam_code(exam_title or "PROVA", exam_date or datetime.now().strftime("%Y-%m-%d"))
    st.code(generated_code, language="text")
    st.caption("Esse codigo sera salvo com a prova e usado no link/QR quando voce quiser gerar consulta automatica.")
    st.info("Os textos do cabecalho e das questoes ja podem ser escritos aqui. No passo seguinte, eu recomendo persistir isso no banco com uma migracao extra.")

    exams = load_exams()
    if st.button("Salvar estrutura da prova", type="primary", use_container_width=True):
        if not exam_title.strip():
            st.warning("Informe um titulo para a prova.")
        else:
            exam = {
                "code": generated_code,
                "title": exam_title.strip(),
                "date": exam_date.strip(),
                "objective_questions": int(objective_questions),
                "essay_questions": int(essay_questions),
                "answer_key": [""] * int(objective_questions),
                "version": exams.get(generated_code, {}).get("version", 0),
                "updated_at": datetime.now().isoformat(timespec="seconds"),
                "header": exam_header,
                "objective_texts": st.session_state.draft_objective_texts[: int(objective_questions)],
                "objective_options": st.session_state.draft_objective_options[: int(objective_questions)],
                "essay_texts": st.session_state.draft_essay_texts[: int(essay_questions)],
            }
            saved_exam = save_exam(exam)
            set_active_exam(saved_exam)
            st.success("Prova criada. O proximo passo e definir o gabarito.")
            go_to_page("2. Definir gabarito")

    selected_exam = render_exam_selector(exams, "Provas ja cadastradas")
    if selected_exam:
        exam_url = build_exam_url(base_url, selected_exam["code"])
        st.markdown(
            f"**Codigo:** `{selected_exam['code']}` | Objetivas: {get_exam_objective_count(selected_exam)} | Dissertativas: {get_exam_essay_count(selected_exam)}"
        )
        if exam_url:
            st.code(exam_url, language="text")
            st.image(build_qr_image(exam_url), caption="QR da prova", width=220)
        else:
            st.info("Se voce preencher a URL publica do app, o QR e o link serao gerados aqui.")

        printable_html = generate_printable_template(selected_exam, exam_url)
        docx_file = build_exam_docx(selected_exam)
        pdf_file = build_exam_pdf(selected_exam)
        export_col1, export_col2, export_col3 = st.columns(3)
        export_col1.download_button(
            "Baixar HTML da prova",
            data=printable_html.encode("utf-8"),
            file_name=f"prova_{selected_exam['code'].lower()}.html",
            mime="text/html",
            use_container_width=True,
        )
        export_col2.download_button(
            "Baixar Word da prova",
            data=docx_file,
            file_name=f"prova_{selected_exam['code'].lower()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        export_col3.download_button(
            "Baixar PDF da prova",
            data=pdf_file,
            file_name=f"prova_{selected_exam['code'].lower()}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )


def render_answer_key_page():
    st.subheader("2. Definir gabarito")
    exams = load_exams()
    exam = st.session_state.active_exam or render_exam_selector(exams, "Escolha a prova para definir o gabarito")
    if exam is None:
        return

    if st.session_state.active_exam_code != exam["code"]:
        set_active_exam(exam)

    total_questions = get_exam_objective_count(exam)
    st.markdown(
        f"**Prova:** {exam['title']} | **Codigo:** `{exam['code']}` | Objetivas: {total_questions} | Dissertativas: {get_exam_essay_count(exam)}"
    )
    render_answer_key_editor(total_questions)

    col1, col2 = st.columns([1.2, 1.2])
    if col1.button("Salvar gabarito desta prova", type="primary", use_container_width=True):
        if not answer_key_complete(st.session_state.draft_answer_key, total_questions):
            st.warning("Preencha todas as questoes objetivas antes de salvar.")
        else:
            saved_exam = exam.copy()
            saved_exam["answer_key"] = st.session_state.draft_answer_key[:total_questions]
            saved_exam["version"] = int(saved_exam.get("version", 0)) + 1
            saved_exam["updated_at"] = datetime.now().isoformat(timespec="seconds")
            persisted_exam = save_exam(saved_exam)
            set_active_exam(persisted_exam)
            st.success("Gabarito salvo. Agora voce pode seguir para a correcao.")
            go_to_page("3. Corrigir gabaritos")

    if col2.button("Limpar escolhas do gabarito", use_container_width=True):
        draft = st.session_state.draft_answer_key.copy()
        for index in range(total_questions):
            draft[index] = ""
        st.session_state.draft_answer_key = draft
        st.info("Escolhas limpas para esta prova.")


def render_active_exam_banner():
    active_exam = st.session_state.active_exam
    if not active_exam:
        st.warning("Cadastre uma prova e salve o gabarito antes de corrigir.")
        return False

    total_questions = get_exam_objective_count(active_exam)
    if not answer_key_complete(active_exam["answer_key"], total_questions):
        st.warning("Esta prova ainda nao tem gabarito completo salvo.")
        return False

    st.info(
        f"Prova ativa: {active_exam['title']} | Codigo: {active_exam['code']} | Objetivas: {total_questions} | Dissertativas: {get_exam_essay_count(active_exam)}"
    )
    return True


def render_correction_page():
    st.subheader("3. Corrigir gabaritos")
    exams = load_exams()
    selected_exam = render_exam_selector(exams, "Escolha a prova para corrigir") if exams else None
    if selected_exam and (st.session_state.active_exam_code != selected_exam["code"]):
        set_active_exam(selected_exam)

    if not render_active_exam_banner():
        render_saved_results(st.session_state.active_exam_code or None)
        return

    active_exam = st.session_state.active_exam
    nome = st.text_input("Nome do aluno", key="nome_aluno_input")
    turma = st.text_input("Turma", key="turma_input")
    foto = st.camera_input("Tire uma foto do gabarito", key="foto_gabarito_input")

    if foto is None:
        st.info("Tire uma foto do gabarito para iniciar a leitura.")
        render_saved_results(active_exam["code"])
        return

    imagem = Image.open(foto)
    imagem_bgr = pil_to_bgr(imagem)

    try:
        resultado = process_answer_sheet(imagem_bgr)
        respostas = resultado["respostas"][: get_exam_objective_count(active_exam)]
        acertos, comparacoes = compute_score(respostas, active_exam["answer_key"])
        st.session_state.last_processed = {
            "nome": nome,
            "turma": turma,
            "respostas": respostas,
            "acertos": acertos,
            "comparacoes": comparacoes,
            "prova": active_exam["code"],
        }

        st.subheader("Resultado da leitura")
        st.write("Prova:", active_exam["title"])
        st.write("Codigo da prova:", active_exam["code"])
        st.write("Aluno:", nome or "Nao informado")
        st.write("Turma:", turma or "Nao informada")
        st.write("Respostas lidas:", " ".join(respostas))
        st.write("Gabarito salvo:", " ".join(active_exam["answer_key"]))
        st.write("Acertos:", acertos)
        st.write("Erros:", len(active_exam["answer_key"]) - acertos)

        comparison_df = pd.DataFrame(comparacoes)
        st.dataframe(comparison_df, use_container_width=True, hide_index=True)

        if st.button("Salvar correcao deste aluno", type="primary", use_container_width=True):
            registro = build_student_record(nome, turma, active_exam, respostas, acertos, comparacoes)
            save_result(registro, active_exam["id"])
            st.success("Correcao salva na turma.")

        render_diagnostics(resultado)
    except ValueError as error:
        st.error(str(error))
        st.image(imagem, caption="Imagem original recebida")

    render_saved_results(active_exam["code"])


def main():
    st.set_page_config(page_title="Leitor de Gabarito", layout="wide")
    init_state()

    exams = load_exams()
    sync_exam_from_query(exams)

    st.title("Leitor de Gabarito")
    st.caption("Fluxo sugerido: criar prova, definir gabarito e depois corrigir a turma.")

    page = st.sidebar.radio(
        "Etapas",
        ["1. Criar prova", "2. Definir gabarito", "3. Corrigir gabaritos"],
        key="sidebar_page",
        index=["1. Criar prova", "2. Definir gabarito", "3. Corrigir gabaritos"].index(st.session_state.nav_page),
    )
    if page != st.session_state.nav_page:
        st.session_state.nav_page = page

    current_page = st.session_state.nav_page

    if current_page == "1. Criar prova":
        render_create_exam_page()
    elif current_page == "2. Definir gabarito":
        render_answer_key_page()
    else:
        render_correction_page()


if __name__ == "__main__":
    main()
